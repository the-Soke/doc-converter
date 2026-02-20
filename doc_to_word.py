"""
doc_to_word.py
==============
Utility that takes the final Markdown (from doc_reconstructor.py) and/or the
extracted JSON (from doc_extractor.py) and generates a professionally formatted
Word (.docx) document.

Font sizes, bolding, headings, bullet points, and emphasis styles from the
original template are all mapped correctly.

Usage
-----
    # From Markdown file (recommended — preserves full structure):
    python doc_to_word.py --markdown-file reconstructed.md --output final.docx

    # From JSON file (uses keys as headings, values as body):
    python doc_to_word.py --json-file result.json --output final.docx

    # From both (Markdown structure + JSON for richer context):
    python doc_to_word.py --markdown-file reconstructed.md --json-file result.json --output final.docx

Requirements
------------
    pip install python-docx

No API key needed — this is 100% local.
"""

from __future__ import annotations

import argparse
import json
import logging
import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor, Inches

# -- Logging ------------------------------------------------------------------
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s  %(levelname)-8s  %(message)s",
    datefmt="%H:%M:%S",
)
log = logging.getLogger(__name__)


# =============================================================================
# 1.  Style configuration  — edit this to match your template's look & feel
# =============================================================================

STYLE_MAP = {
    # heading level : (font_size_pt, bold, italic, color_hex, space_before_pt, space_after_pt, alignment)
    "h1": (22, True,  False, "1F3864", 18, 6,  WD_ALIGN_PARAGRAPH.LEFT),
    "h2": (16, True,  False, "2E5496", 14, 4,  WD_ALIGN_PARAGRAPH.LEFT),
    "h3": (13, True,  False, "2E5496", 10, 3,  WD_ALIGN_PARAGRAPH.LEFT),
    "h4": (12, True,  True,  "404040", 8,  2,  WD_ALIGN_PARAGRAPH.LEFT),
    "h5": (11, False, True,  "404040", 6,  2,  WD_ALIGN_PARAGRAPH.LEFT),
    "h6": (11, False, True,  "666666", 4,  2,  WD_ALIGN_PARAGRAPH.LEFT),
    # body text
    "body":   (11, False, False, "000000", 0, 4,  WD_ALIGN_PARAGRAPH.LEFT),
    # bullet list item
    "bullet": (11, False, False, "000000", 0, 2,  WD_ALIGN_PARAGRAPH.LEFT),
    # numbered list item
    "number": (11, False, False, "000000", 0, 2,  WD_ALIGN_PARAGRAPH.LEFT),
    # horizontal rule — rendered as a bottom-border paragraph
    "hr":     (6,  False, False, "CCCCCC", 4, 4,  WD_ALIGN_PARAGRAPH.LEFT),
}

FONT_NAME = "Calibri"            # change to "Arial", "Times New Roman", etc.
PAGE_WIDTH_INCHES  = 8.5
PAGE_HEIGHT_INCHES = 11.0
MARGIN_INCHES      = 1.0


# =============================================================================
# 2.  Low-level python-docx helpers
# =============================================================================

def _set_paragraph_style(para, style_key: str) -> None:
    """Apply font size, bold, italic, color and spacing from STYLE_MAP."""
    size, bold, italic, color_hex, space_before, space_after, alignment = STYLE_MAP[style_key]

    para.alignment = alignment
    para.paragraph_format.space_before = Pt(space_before)
    para.paragraph_format.space_after  = Pt(space_after)

    for run in para.runs:
        run.font.name    = FONT_NAME
        run.font.size    = Pt(size)
        run.font.bold    = bold
        run.font.italic  = italic
        run.font.color.rgb = RGBColor(
            int(color_hex[0:2], 16),
            int(color_hex[2:4], 16),
            int(color_hex[4:6], 16),
        )


def _add_horizontal_rule(doc: Document) -> None:
    """Add a thin horizontal line (bottom border on an empty paragraph)."""
    para = doc.add_paragraph()
    pPr  = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"),   "single")
    bottom.set(qn("w:sz"),    "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "AAAAAA")
    pBdr.append(bottom)
    pPr.append(pBdr)
    para.paragraph_format.space_before = Pt(STYLE_MAP["hr"][4])
    para.paragraph_format.space_after  = Pt(STYLE_MAP["hr"][5])


def _add_bullet_paragraph(doc: Document, text: str, numbered: bool = False) -> None:
    """Add a native Word list paragraph (no unicode bullet characters)."""
    style = "List Number" if numbered else "List Bullet"
    try:
        para = doc.add_paragraph(style=style)
    except KeyError:
        # Fallback if built-in style isn't present
        para = doc.add_paragraph()

    _apply_inline_formatting(para, text)
    _set_paragraph_style(para, "number" if numbered else "bullet")


def _apply_inline_formatting(para, text: str) -> None:
    """
    Parse inline Markdown formatting within a line and add styled runs.
    Handles: **bold**, *italic*, ***bold+italic***, `code`, and plain text.
    """
    # Pattern captures groups: bold+italic / bold / italic / code / plain
    pattern = re.compile(
        r"(\*\*\*(.+?)\*\*\*)"    # bold + italic
        r"|(\*\*(.+?)\*\*)"       # bold
        r"|(\*(.+?)\*)"           # italic
        r"|(`(.+?)`)"             # inline code
        r"|([^*`]+)"              # plain text
    )

    for match in pattern.finditer(text):
        if match.group(1):    # bold + italic
            run = para.add_run(match.group(2))
            run.bold = True
            run.italic = True
        elif match.group(3):  # bold
            run = para.add_run(match.group(4))
            run.bold = True
        elif match.group(5):  # italic
            run = para.add_run(match.group(6))
            run.italic = True
        elif match.group(7):  # inline code
            run = para.add_run(match.group(8))
            run.font.name = "Courier New"
            run.font.size = Pt(10)
        else:                  # plain text
            para.add_run(match.group(0))


def _set_page_size(doc: Document) -> None:
    """Set US Letter page size and margins."""
    section = doc.sections[0]
    section.page_width    = Inches(PAGE_WIDTH_INCHES)
    section.page_height   = Inches(PAGE_HEIGHT_INCHES)
    section.left_margin   = Inches(MARGIN_INCHES)
    section.right_margin  = Inches(MARGIN_INCHES)
    section.top_margin    = Inches(MARGIN_INCHES)
    section.bottom_margin = Inches(MARGIN_INCHES)


# =============================================================================
# 3.  Markdown → Word converter
# =============================================================================

def markdown_to_word(markdown_text: str, doc: Document) -> None:
    """
    Parse a Markdown string and write it into a python-docx Document.

    Supported Markdown elements:
      - # H1 through ###### H6
      - **bold**, *italic*, ***bold+italic***, `code`
      - - / * unordered lists  (nested with spaces)
      - 1. ordered lists
      - --- / *** horizontal rules
      - Blank lines (paragraph breaks)
      - Plain body paragraphs
    """
    lines = markdown_text.splitlines()
    i = 0
    list_counter: dict[int, int] = {}  # indent_level -> current number

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # ── Blank line ──────────────────────────────────────────────────────
        if not stripped:
            i += 1
            continue

        # ── Horizontal rule ─────────────────────────────────────────────────
        if re.fullmatch(r"[-*_]{3,}", stripped):
            _add_horizontal_rule(doc)
            i += 1
            continue

        # ── ATX Headings (# H1 … ###### H6) ────────────────────────────────
        heading_match = re.match(r"^(#{1,6})\s+(.*)", stripped)
        if heading_match:
            level  = len(heading_match.group(1))
            text   = heading_match.group(2).strip()
            key    = f"h{level}"
            size, bold, italic, color_hex, sb, sa, align = STYLE_MAP[key]

            # Map to Word heading styles
            heading_style_map = {
                1: "Heading 1", 2: "Heading 2", 3: "Heading 3",
                4: "Heading 4", 5: "Heading 5", 6: "Heading 6",
            }
            try:
                para = doc.add_paragraph(style=heading_style_map[level])
            except KeyError:
                para = doc.add_paragraph()

            _apply_inline_formatting(para, text)
            _set_paragraph_style(para, key)
            i += 1
            continue

        # ── Unordered list items (-, *, +) ──────────────────────────────────
        ul_match = re.match(r"^(\s*)([-*+])\s+(.*)", line)
        if ul_match:
            text = ul_match.group(3)
            _add_bullet_paragraph(doc, text, numbered=False)
            i += 1
            continue

        # ── Ordered list items (1. 2. etc.) ─────────────────────────────────
        ol_match = re.match(r"^(\s*)\d+[.)]\s+(.*)", line)
        if ol_match:
            text = ol_match.group(2)
            _add_bullet_paragraph(doc, text, numbered=True)
            i += 1
            continue

        # ── Setext-style H1 (underlined with ===) ───────────────────────────
        if i + 1 < len(lines) and re.fullmatch(r"=+", lines[i + 1].strip()):
            para = doc.add_paragraph(style="Heading 1")
            _apply_inline_formatting(para, stripped)
            _set_paragraph_style(para, "h1")
            i += 2
            continue

        # ── Setext-style H2 (underlined with ---) ───────────────────────────
        if i + 1 < len(lines) and re.fullmatch(r"-+", lines[i + 1].strip()) and len(lines[i + 1].strip()) > 2:
            para = doc.add_paragraph(style="Heading 2")
            _apply_inline_formatting(para, stripped)
            _set_paragraph_style(para, "h2")
            i += 2
            continue

        # ── Body paragraph ───────────────────────────────────────────────────
        para = doc.add_paragraph()
        _apply_inline_formatting(para, stripped)
        _set_paragraph_style(para, "body")
        i += 1


# =============================================================================
# 4.  JSON → Word converter  (fallback / standalone)
# =============================================================================

def json_to_word(data: dict, doc: Document) -> None:
    """
    Convert extracted JSON (key → value pairs) into a Word document.
    Each key becomes an H2 section heading; the value becomes body text.
    Null values are shown as '[Not provided]'.
    """
    for key, value in data.items():
        # Format key as a readable heading
        heading_text = key.replace("_", " ").title()
        para = doc.add_paragraph(style="Heading 2")
        para.add_run(heading_text)
        _set_paragraph_style(para, "h2")

        # Value
        content = str(value) if value is not None else "[Not provided]"

        # If value contains newlines, split into separate paragraphs
        for line in content.splitlines():
            line = line.strip()
            if not line:
                continue
            # Detect bullet-like lines inside values
            if re.match(r"^[-*+]\s+", line):
                _add_bullet_paragraph(doc, re.sub(r"^[-*+]\s+", "", line))
            else:
                body_para = doc.add_paragraph()
                _apply_inline_formatting(body_para, line)
                _set_paragraph_style(body_para, "body")

        # Visual separator between sections
        _add_horizontal_rule(doc)


# =============================================================================
# 5.  Main build function
# =============================================================================

def build_word_document(
    markdown_text: str | None = None,
    json_data: dict | None = None,
    output_path: str | Path = "final.docx",
) -> Path:
    """
    Create a professionally formatted Word document from Markdown and/or JSON.

    Priority:
      - If markdown_text is provided, it drives the full structure.
      - If only json_data is provided, a structured section-per-key layout is used.
      - If both are provided, Markdown is used as the primary structure.

    Parameters
    ----------
    markdown_text : Markdown string from doc_reconstructor.py (optional)
    json_data     : dict from doc_extractor.py result.json (optional)
    output_path   : where to save the .docx file

    Returns
    -------
    Path to the saved .docx file
    """
    if not markdown_text and not json_data:
        raise ValueError("Provide at least one of: markdown_text or json_data.")

    doc = Document()
    _set_page_size(doc)

    # ── Override default Normal style font ───────────────────────────────────
    normal_style          = doc.styles["Normal"]
    normal_style.font.name = FONT_NAME
    normal_style.font.size = Pt(STYLE_MAP["body"][0])

    # ── Build content ────────────────────────────────────────────────────────
    if markdown_text:
        log.info("Converting Markdown to Word ...")
        markdown_to_word(markdown_text, doc)

        # Optionally append JSON section at the end if both provided
        if json_data:
            doc.add_page_break()
            title_para = doc.add_paragraph(style="Heading 1")
            title_para.add_run("Extracted Data Reference")
            _set_paragraph_style(title_para, "h1")
            json_to_word(json_data, doc)
    else:
        log.info("Converting JSON to Word ...")
        json_to_word(json_data, doc)  # type: ignore[arg-type]

    # ── Save ─────────────────────────────────────────────────────────────────
    out = Path(output_path)
    doc.save(str(out))
    log.info("Word document saved to '%s'.", out.resolve())
    return out


# =============================================================================
# 6.  CLI entry-point
# =============================================================================

def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Generate a professionally formatted Word document from "
                    "Markdown and/or JSON using python-docx.",
        formatter_class=argparse.RawDescriptionHelpFormatter,
    )
    parser.add_argument(
        "--markdown-file", default=None,
        help="Path to reconstructed.md from doc_reconstructor.py.",
    )
    parser.add_argument(
        "--json-file", default=None,
        help="Path to result.json from doc_extractor.py.",
    )
    parser.add_argument(
        "--output", default="final.docx",
        help="Output .docx filename (default: final.docx).",
    )
    return parser.parse_args()


def main() -> None:
    args = parse_args()

    if not args.markdown_file and not args.json_file:
        print("ERROR: Provide --markdown-file and/or --json-file.")
        sys.exit(1)

    # Load Markdown
    markdown_text = None
    if args.markdown_file:
        md_path = Path(args.markdown_file)
        if not md_path.exists():
            print(f"ERROR: Markdown file not found: {md_path}")
            sys.exit(1)
        markdown_text = md_path.read_text(encoding="utf-8")
        log.info("Loaded Markdown from '%s' (%d chars).", md_path.name, len(markdown_text))

    # Load JSON
    json_data = None
    if args.json_file:
        json_path = Path(args.json_file)
        if not json_path.exists():
            print(f"ERROR: JSON file not found: {json_path}")
            sys.exit(1)
        json_data = json.loads(json_path.read_text(encoding="utf-8"))
        log.info("Loaded JSON from '%s' (%d fields).", json_path.name, len(json_data))

    build_word_document(
        markdown_text=markdown_text,
        json_data=json_data,
        output_path=args.output,
    )

    log.info("Done! Open '%s' in Word or VSCode to review.", args.output)


if __name__ == "__main__":
    main()
